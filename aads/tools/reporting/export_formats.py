"""
AUDAS Report Exporters — Converts Markdown reports to .docx (Word) and .pdf (PDF) formats.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List


def _clean_inline_markdown(text: str) -> str:
    """Strip basic markdown links and format bold/italic for plain text."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def export_markdown_to_docx(md_content: str, output_path: Path) -> Path:
    """
    Converts a Markdown string into a styled Microsoft Word (.docx) document.
    """
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    doc = docx.Document()

    # Set page margins to 0.75 in
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    lines = md_content.splitlines()
    i = 0
    in_code_block = False
    table_lines: List[str] = []

    def _flush_table(tbl_lines: List[str]):
        if not tbl_lines:
            return
        rows = []
        for line in tbl_lines:
            stripped = line.strip()
            if stripped.startswith("|") and stripped.endswith("|"):
                cells = [c.strip() for c in stripped[1:-1].split("|")]
                if all(re.match(r"^:?-+:?$", c) for c in cells if c):
                    continue
                rows.append(cells)

        if not rows:
            return

        cols_count = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=cols_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx]
            for c_idx in range(cols_count):
                cell_value = row_data[c_idx] if c_idx < len(row_data) else ""
                cell = row.cells[c_idx]
                cell.text = _clean_inline_markdown(cell_value)
                
                # Style header vs data
                if r_idx == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2D1B69"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(9.5)
                else:
                    if r_idx % 2 == 1:
                        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>')
                        cell._tc.get_or_add_tcPr().append(shd)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9.5)

        doc.add_paragraph()  # spacing

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle Table Gathering
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines.append(stripped)
            i += 1
            continue
        elif table_lines:
            _flush_table(table_lines)
            table_lines = []

        # Handle Code Blocks
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(71, 85, 105)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            h = doc.add_heading(level=1)
            run = h.add_run(stripped[2:].strip())
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(45, 27, 105)
        elif stripped.startswith("## "):
            h = doc.add_heading(level=2)
            run = h.add_run(stripped[3:].strip())
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(79, 70, 229)
        elif stripped.startswith("### "):
            h = doc.add_heading(level=3)
            run = h.add_run(stripped[4:].strip())
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(15, 23, 42)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(p, bullet_text, Pt(10))
        elif re.match(r"^\d+\.\s", stripped):
            num_text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_runs(p, num_text, Pt(10))
        elif stripped.startswith(">"):
            quote_text = stripped.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(f"▎ {quote_text}")
            run.font.italic = True
            run.font.color.rgb = RGBColor(99, 102, 241)
            run.font.size = Pt(10)
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, stripped, Pt(10.5))

        i += 1

    if table_lines:
        _flush_table(table_lines)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _add_formatted_runs(paragraph, text: str, font_size):
    """Parses bold **text** and code `code` in text and adds runs."""
    from docx.shared import Pt, RGBColor

    tokens = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.font.bold = True
            run.font.size = font_size
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(147, 51, 234)
        else:
            run = paragraph.add_run(tok)
            run.font.size = font_size


def export_markdown_to_pdf(md_content: str, output_path: Path, title: str = "AUDAS Executive Summary") -> Path:
    """
    Converts a Markdown string into a styled PDF using ReportLab.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        HRFlowable,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=40,
        rightMargin=40,
        topMargin=40,
        bottomMargin=40,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#1e1b4b"),
        spaceAfter=6,
    )
    h2_style = ParagraphStyle(
        "DocH2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=17,
        textColor=colors.HexColor("#4338ca"),
        spaceBefore=12,
        spaceAfter=5,
        keepWithNext=True,
    )
    h3_style = ParagraphStyle(
        "DocH3",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#0f172a"),
        spaceBefore=8,
        spaceAfter=4,
        keepWithNext=True,
    )
    body_style = ParagraphStyle(
        "DocBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor("#1e293b"),
        spaceAfter=4,
    )
    bullet_style = ParagraphStyle(
        "DocBullet",
        parent=body_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=3,
    )
    quote_style = ParagraphStyle(
        "DocQuote",
        parent=body_style,
        fontName="Helvetica-Oblique",
        textColor=colors.HexColor("#6366f1"),
        leftIndent=15,
        spaceBefore=4,
        spaceAfter=4,
    )
    table_cell_style = ParagraphStyle(
        "TableCell",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#0f172a"),
    )
    table_header_style = ParagraphStyle(
        "TableHeader",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=8.5,
        leading=11,
        textColor=colors.white,
    )

    story = []
    lines = md_content.splitlines()
    i = 0
    in_code_block = False
    table_lines: List[str] = []

    def _flush_pdf_table(tbl_lines: List[str]):
        if not tbl_lines:
            return
        rows = []
        for line in tbl_lines:
            stripped = line.strip()
            if stripped.startswith("|") and stripped.endswith("|"):
                cells = [c.strip() for c in stripped[1:-1].split("|")]
                if all(re.match(r"^:?-+:?$", c) for c in cells if c):
                    continue
                rows.append(cells)

        if not rows:
            return

        cols_count = max(len(r) for r in rows)
        table_data = []
        for r_idx, row_data in enumerate(rows):
            row_cells = []
            for c_idx in range(cols_count):
                val = row_data[c_idx] if c_idx < len(row_data) else ""
                val_xml = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", val)
                val_xml = re.sub(r"`([^`]+)`", r"<font face='Courier'>\1</font>", val_xml)
                st = table_header_style if r_idx == 0 else table_cell_style
                row_cells.append(Paragraph(val_xml, st))
            table_data.append(row_cells)

        col_width = 532 / cols_count if cols_count else 100
        t = Table(table_data, colWidths=[col_width] * cols_count)
        t.setStyle(
            TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#312e81")),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#ffffff"), colors.HexColor("#f8fafc")]),
            ])
        )
        story.append(Spacer(1, 4))
        story.append(t)
        story.append(Spacer(1, 6))

    def _safe_p(txt: str, st: ParagraphStyle) -> Paragraph:
        try:
            return Paragraph(txt, st)
        except Exception:
            # Fallback: strip XML tags to prevent formatting exceptions
            clean = re.sub(r"<[^>]+>", "", txt)
            clean = clean.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
            clean = clean.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            return Paragraph(clean, st)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle Tables
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines.append(stripped)
            i += 1
            continue
        elif table_lines:
            _flush_pdf_table(table_lines)
            table_lines = []

        # Handle Code Blocks
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            code_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(_safe_p(f"<font face='Courier' size=8 color='#475569'>{code_line}</font>", body_style))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        formatted = stripped
        formatted = formatted.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        formatted = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", formatted)
        formatted = re.sub(r"\*([^*]+)\*", r"<i>\1</i>", formatted)
        formatted = re.sub(r"`([^`]+)`", r"<font face='Courier' color='#7c3aed'>\1</font>", formatted)
        formatted = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"<b>\1</b>", formatted)

        # Headings
        if stripped.startswith("# "):
            text = formatted[2:].strip()
            story.append(_safe_p(text, title_style))
            story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#6366f1"), spaceAfter=10))
        elif stripped.startswith("## "):
            text = formatted[3:].strip()
            story.append(_safe_p(text, h2_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1"), spaceAfter=6))
        elif stripped.startswith("### "):
            text = formatted[4:].strip()
            story.append(_safe_p(text, h3_style))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = formatted[2:].strip()
            story.append(_safe_p(f"• {text}", bullet_style))
        elif re.match(r"^\d+\.\s", stripped):
            num = re.match(r"^(\d+\.)\s", stripped).group(1)
            text = re.sub(r"^\d+\.\s", "", formatted)
            story.append(_safe_p(f"<b>{num}</b> {text}", bullet_style))
        elif stripped.startswith(">"):
            text = formatted.lstrip("> ").strip()
            story.append(_safe_p(f"▎ {text}", quote_style))
        else:
            story.append(_safe_p(formatted, body_style))

        i += 1

    if table_lines:
        _flush_pdf_table(table_lines)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)
    return output_path
